"""
Generate .docx files from structured editor content using python-docx (FREE).

Preserves paragraphs, alignment, and inline formatting (bold, italic, underline, size).
"""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.services.document.font_mapper import resolve_font


def build_docx(payload: dict) -> bytes:
    """
    Build a Word document from editor JSON.

    Expected payload:
    {
      "fontFamily": "mangal",
      "blocks": [
        {
          "type": "paragraph",
          "align": "left" | "center" | "right",
          "runs": [
            {
              "text": "...",
              "bold": false,
              "italic": false,
              "underline": false,
              "fontSize": 14,
              "uncertain": false
            }
          ]
        }
      ]
    }
    """
    doc = Document()
    font_name = resolve_font(payload.get("fontFamily"))

    for block in payload.get("blocks", []):
        paragraph = doc.add_paragraph()
        align = block.get("align", "left")
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for run_data in block.get("runs", []):
            text = run_data.get("text", "")
            if not text:
                continue
            run = paragraph.add_run(text)
            run.bold = bool(run_data.get("bold"))
            run.italic = bool(run_data.get("italic"))
            run.underline = bool(run_data.get("underline"))
            run.font.name = font_name
            run.font.size = Pt(int(run_data.get("fontSize", 14)))
            if run_data.get("uncertain"):
                run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
